#!/usr/bin/env python3
"""
Zero Crisis MCP server.
Exposes one tool: create_donation_receipt
Generates a PDF and returns a time-limited download link.
"""

import os
import re
import uuid
import subprocess
from datetime import datetime
from pathlib import Path

from fastmcp import FastMCP
from starlette.responses import FileResponse, JSONResponse
from starlette.routing import Route

from generate_receipts import (
    TEMPLATE_PATH, ANCHORS,
    replace_in_document, _fix_donor_section, _fix_date_font, _fix_image_wrap,
    fmt_amount, fmt_date,
)
from docx import Document

TEMP_DIR = Path("/tmp/receipts")
TEMP_DIR.mkdir(parents=True, exist_ok=True)

mcp = FastMCP("Zero Crisis Receipt Generator")


def base_url() -> str:
    domain = os.environ.get("RAILWAY_PUBLIC_DOMAIN", "localhost:8000")
    scheme = "https" if "localhost" not in domain else "http"
    return f"{scheme}://{domain}"


@mcp.tool()
def create_donation_receipt(
    donor_name: str,
    donor_email: str,
    amount: float,
    processor: str,
    fee: float = 0.0,
    date: str = "",
) -> str:
    """
    Generate a PDF donation receipt and return a download link.

    Args:
        donor_name:  Full name of the donor (e.g. "Jane Smith")
        donor_email: Donor's email address
        amount:      Gross donation amount as a number (e.g. 50.0)
        processor:   Payment processor — "Venmo", "PayPal", or "Zelle"
        fee:         Transaction fee (default 0.0)
        date:        Donation date as MM/DD/YYYY (defaults to today)
    """
    if processor not in ("Venmo", "PayPal", "Zelle"):
        return f'Invalid processor "{processor}". Use Venmo, PayPal, or Zelle.'

    if not date:
        date = datetime.today().strftime("%m/%d/%Y")

    date_fmt = fmt_date(date)
    net      = amount - abs(fee)
    doc_id   = str(int(datetime.now().timestamp()))

    replacements = {
        ANCHORS["donor_name"]:     donor_name,
        ANCHORS["date"]:           date_fmt,
        ANCHORS["gross_amount"]:   fmt_amount(amount),
        ANCHORS["fee"]:            fmt_amount(abs(fee)),
        ANCHORS["net_amount"]:     fmt_amount(net),
        ANCHORS["transaction_id"]: doc_id,
        ANCHORS["processor"]:      processor,
    }
    replacements = {k: v for k, v in replacements.items() if k and v}

    doc = Document(TEMPLATE_PATH)
    replace_in_document(doc, replacements)
    _fix_donor_section(doc, donor_email)
    _fix_date_font(doc)
    _fix_image_wrap(doc)

    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[5]
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.startswith("$"):
                        run.text = "-" + run.text

    uid       = uuid.uuid4().hex[:8]
    safe_name = donor_name.replace(" ", "_").replace("/", "-")
    temp_docx = TEMP_DIR / f"{uid}_temp.docx"
    pdf_name  = f"{uid}_{doc_id}_Donation_Receipt_{safe_name}.pdf"
    output_pdf = TEMP_DIR / pdf_name

    doc.save(str(temp_docx))

    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(TEMP_DIR), str(temp_docx)],
        capture_output=True, text=True,
    )
    temp_docx.unlink(missing_ok=True)

    lo_output = TEMP_DIR / f"{uid}_temp.pdf"
    if lo_output.exists():
        lo_output.rename(output_pdf)
    else:
        return f"ERROR: PDF generation failed. {result.stderr.strip()}"

    download_url = f"{base_url()}/receipt/{pdf_name}"
    return (
        f"Receipt created for {donor_name} ({date_fmt}, {processor}, "
        f"${amount:.2f}).\n\nDownload: {download_url}"
    )


# ── File serving route ────────────────────────────────────────────────────────

async def serve_receipt(request):
    filename = request.path_params["filename"]
    if not re.match(r'^[\w\-\.]+\.pdf$', filename):
        return JSONResponse({"error": "Invalid filename"}, status_code=400)
    path = TEMP_DIR / filename
    if not path.exists():
        return JSONResponse({"error": "File not found or expired"}, status_code=404)
    return FileResponse(str(path), media_type="application/pdf", filename=filename)


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    from starlette.applications import Starlette
    from starlette.routing import Mount

    mcp_app = mcp.http_app(transport="sse")

    app = Starlette(routes=[
        Route("/receipt/{filename}", serve_receipt),
        Mount("/", app=mcp_app),
    ])

    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
