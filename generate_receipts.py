#!/usr/bin/env python3
"""
Donation receipt generator.

Reads a CSV export (Venmo, PayPal, Zelle, etc.) and generates one PDF receipt
per row using the .docx template.

── Quick-start ──────────────────────────────────────────────────────────────────
  python generate_receipts.py donations.csv

── Config to update ─────────────────────────────────────────────────────────────
  1. COLUMNS      → map to your CSV's actual column headers
  2. PAYMENT_PROCESSOR → "Venmo", "PayPal", "Zelle", etc.
  3. HAS_FEE      → set False if your processor doesn't charge fees
"""

import csv
import os
import sys
from datetime import datetime
from docx import Document

# ── Configuration ─────────────────────────────────────────────────────────────

TEMPLATE_PATH     = "PayPal_Donation Receipt.docx"
OUTPUT_DIR        = "receipts"
DOC_ID_START = 26290660571111099

# ── Per-processor config ───────────────────────────────────────────────────────
# Each entry: column mapping + processor name + whether fees apply.
# The script picks the right one based on the CSV filename.

VENMO_CONFIG = {
    "processor": "Venmo",
    "has_fee": True,
    "columns": {
        "donor_name":   "From",
        "donor_email":  "Donor email",
        "date":         "Date",
        "gross_amount": "Amount (total)",
        "fee":          "Amount (fee)",
        "net_amount":   "Amount (net)",
    },
}

PAYPAL_CONFIG = {
    "processor": "PayPal",
    "has_fee": True,
    "fee_is_negative": True,   # PayPal exports fee as a negative number
    "columns": {
        "donor_name":   "Name",
        "donor_email":  "From Email Address",
        "date":         "Date",
        "gross_amount": "Gross",
        "fee":          "Fee",
        "net_amount":   "Net",
    },
}

ZELLE_CONFIG = {
    "processor": "Zelle",
    "has_fee": False,
    # Zelle puts the donor name inside the DESCRIPTION field.
    # Pattern captures everything between "ZELLE FROM " and " ON \d".
    "name_from_description": r"ZELLE FROM (.+?) ON \d",
    "columns": {
        "donor_name":   "",             # parsed from description instead
        "donor_email":  "",             # Zelle has no email
        "date":         "DATE",
        "gross_amount": "AMOUNT",
        "fee":          "",
        "net_amount":   "",
        "description":  "DESCRIPTION",  # used for name extraction
    },
}

# Filename-based lookup (lowercase filename → config dict)
PROCESSOR_CONFIGS = {
    "venmo.csv":  VENMO_CONFIG,
    "paypal.csv": PAYPAL_CONFIG,
    "zelle.csv":  ZELLE_CONFIG,
}

# ── Template anchor values (what's currently in the .docx sample) ─────────────
# These are what gets replaced. Only change if you update the template file.

ANCHORS = {
    "donor_name":     "NAME",
    "donor_email":    "EMAIL",
    "date":           "DATE",
    "gross_amount":   "GROSS",
    "fee":            "FEE",
    "net_amount":     "DEDUCTIBLE",
    "transaction_id": "26290660571111021",
    "processor":      "PayPal",
}

# ── Text replacement (handles values split across multiple Word runs) ──────────

def _replace_in_paragraph(paragraph, replacements: dict):
    """Replace text in a paragraph, merging split runs when needed."""
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(old in full_text for old in replacements):
        return

    # Build a character → run-index map based on current run lengths
    char_to_run = []
    for idx, run in enumerate(paragraph.runs):
        char_to_run.extend([idx] * len(run.text))

    # Process longest keys first to avoid substring collisions
    for old in sorted(replacements, key=len, reverse=True):
        new = replacements[old]
        pos = full_text.find(old)
        if pos == -1:
            continue
        end = pos + len(old)

        affected = sorted(set(char_to_run[pos:end]))
        first_run = paragraph.runs[affected[0]]

        # Offset of first affected run in full_text
        run_offset = sum(len(paragraph.runs[i].text) for i in range(affected[0]))
        # Offset of end of last affected run in full_text
        last_run_end = sum(len(paragraph.runs[i].text) for i in range(affected[-1] + 1))

        before = full_text[run_offset:pos]
        after  = full_text[end:last_run_end]

        first_run.text = before + new + after
        for idx in affected[1:]:
            paragraph.runs[idx].text = ""

        # Rebuild for next iteration
        full_text = full_text[:pos] + new + full_text[end:]
        char_to_run = []
        for idx, run in enumerate(paragraph.runs):
            char_to_run.extend([idx] * len(run.text))


def replace_in_document(doc: Document, replacements: dict):
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)


def _fix_date_font(doc: Document):
    """Clear the explicit font on the date cell so it inherits the document default."""
    for table in doc.tables:
        for row in table.rows:
            date_cell = row.cells[0]
            for para in date_cell.paragraphs:
                for run in para.runs:
                    run.font.name = None


def _fix_image_wrap(doc: Document):
    """Replace wrapThrough with wrapNone so the logo floats freely at its
    position without splitting text or reserving paragraph height."""
    from lxml import etree
    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for para in doc.paragraphs:
        for anchor in para._element.findall(".//{%s}anchor" % WP):
            wrap_through = anchor.find("{%s}wrapThrough" % WP)
            if wrap_through is not None:
                idx = list(anchor).index(wrap_through)
                anchor.remove(wrap_through)
                anchor.insert(idx, etree.Element("{%s}wrapNone" % WP))


def _fix_donor_section(doc: Document, donor_email: str):
    """
    Two fixes for Para 1 (Donor / Donee header):
      1. Insert the donor email into the Email: run, which has no placeholder.
      2. Collapse all multi-tab separators down to one tab and add an explicit
         tab stop at half the content width (4680 twips) so the Donee column
         never shifts regardless of donor name / email length.
    """
    from lxml import etree
    W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    XML = "http://www.w3.org/XML/1998/namespace"

    para = doc.paragraphs[1]
    r_elems = para._p.findall("{%s}r" % W)

    def keep_one_tab(r_elem):
        """Remove all but the first <w:tab/> in a run element."""
        tabs = r_elem.findall("{%s}tab" % W)
        for tab in tabs[1:]:
            r_elem.remove(tab)

    # ── (Donor) row: run[0] has (Donor) + 7 tabs ────────────────────────────
    keep_one_tab(r_elems[0])

    # ── Name row: run[4] has " NAME" + 6 tabs + "Name: " ────────────────────
    keep_one_tab(r_elems[4])

    # ── Email row: run[6] has <br/> + "Email:" + 7 tabs, no placeholder ─────
    r6 = r_elems[6]
    t_elem = r6.find("{%s}t" % W)
    if t_elem is not None:
        t_elem.text = f"Email: {donor_email}"
        t_elem.set("{%s}space" % XML, "preserve")
    keep_one_tab(r6)

    # ── Donee-only rows: run[8] (<br/> + 8 tabs + "Email: ") ────────────────
    #    and run[9] (<br/> + 8 tabs + "Website: …")
    if len(r_elems) > 8:
        keep_one_tab(r_elems[8])
    if len(r_elems) > 9:
        keep_one_tab(r_elems[9])

    # ── Add explicit tab stop at 4680 twips (≈ half page content width) ─────
    pPr = para._p.find("{%s}pPr" % W)
    if pPr is None:
        pPr = etree.SubElement(para._p, "{%s}pPr" % W)
    tabs_elem = pPr.find("{%s}tabs" % W)
    if tabs_elem is not None:
        pPr.remove(tabs_elem)
    tabs_elem = etree.SubElement(pPr, "{%s}tabs" % W)
    tab_stop = etree.SubElement(tabs_elem, "{%s}tab" % W)
    tab_stop.set("{%s}val" % W, "left")
    tab_stop.set("{%s}pos" % W, "4680")

# ── Formatting helpers ─────────────────────────────────────────────────────────

def parse_amount(value: str) -> float:
    return float(value.replace("$", "").replace(",", "").strip())


def fmt_amount(value: float) -> str:
    return f"${value:.2f}"


def fmt_date(raw: str) -> str:
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y"):
        try:
            return datetime.strptime(raw.strip(), fmt).strftime("%B %d, %Y")
        except ValueError:
            continue
    return raw


def safe_filename(donor_name: str, date: str, doc_id: str = "") -> str:
    safe_name = donor_name.replace("/", "-")
    if doc_id:
        return f"{doc_id}_Donation Receipt_{safe_name}.pdf"
    # Fallback (no doc_id) — used only by the manifest builder
    name = donor_name.replace(" ", "_").replace("/", "-")
    date = date.replace(" ", "_").replace(",", "").replace("/", "-")
    return f"receipt_{name}_{date}.pdf"

# ── Main ───────────────────────────────────────────────────────────────────────

def generate_receipts(csv_path: str):
    if not os.path.exists(TEMPLATE_PATH):
        print(f"ERROR: Receipt template not found: '{TEMPLATE_PATH}'")
        print("Make sure 'PayPal_Donation Receipt.docx' is in the same folder as this script.")
        sys.exit(1)

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    filename = os.path.basename(csv_path).lower()
    config   = PROCESSOR_CONFIGS.get(filename)
    if config is None:
        known = ", ".join(PROCESSOR_CONFIGS)
        print(f"Unknown CSV filename '{filename}'. Expected one of: {known}")
        sys.exit(1)

    COLUMNS           = config["columns"]
    PAYMENT_PROCESSOR = config["processor"]
    HAS_FEE           = config["has_fee"]
    FEE_IS_NEGATIVE   = config.get("fee_is_negative", False)
    NAME_PATTERN      = config.get("name_from_description")

    rows = None
    for enc in ("utf-8-sig", "cp1252", "latin-1"):
        try:
            with open(csv_path, newline="", encoding=enc) as f:
                rows = list(csv.DictReader(f))
            break
        except UnicodeDecodeError:
            continue
    if rows is None:
        print(f"ERROR: Could not read '{csv_path}' — try re-saving it as UTF-8.")
        sys.exit(1)

    print(f"Detected: {PAYMENT_PROCESSOR} | {len(rows)} rows → '{OUTPUT_DIR}/'...")

    seen_filenames: dict[str, int] = {}
    doc_counter = 0
    manifest: list[dict] = []

    for i, row in enumerate(rows, 1):
        def col(key):
            col_name = COLUMNS.get(key, "")
            return row.get(col_name, "").strip() if col_name else ""

        # For Zelle: parse donor name out of the description field
        if NAME_PATTERN:
            import re
            desc  = col("description")
            match = re.search(NAME_PATTERN, desc)
            if not match:
                print(f"  [row {i}] skipped (name not found in description)")
                continue
            donor_name_raw = match.group(1).title()
        else:
            donor_name_raw = col("donor_name")

        # Skip footer / summary rows with no donor name or amount
        if not donor_name_raw or not col("gross_amount"):
            print(f"  [row {i}] skipped (empty)")
            continue

        doc_counter += 1
        donor_name  = donor_name_raw
        donor_email = col("donor_email")
        date_fmt    = fmt_date(col("date"))
        gross       = parse_amount(col("gross_amount"))
        fee_raw     = parse_amount(col("fee")) if HAS_FEE and col("fee") else 0.0
        fee         = abs(fee_raw) if FEE_IS_NEGATIVE else fee_raw
        net_raw     = col("net_amount")
        net         = parse_amount(net_raw) if net_raw else gross - fee
        doc_id      = str(DOC_ID_START + doc_counter - 1)

        gross_str = fmt_amount(gross)
        fee_str   = fmt_amount(fee)
        net_str   = fmt_amount(net)

        replacements = {
            ANCHORS["donor_name"]:     donor_name,
            ANCHORS["date"]:           date_fmt,
            ANCHORS["gross_amount"]:   gross_str,
            ANCHORS["fee"]:            fee_str,
            ANCHORS["net_amount"]:     net_str,
            ANCHORS["transaction_id"]: doc_id,
            ANCHORS["processor"]:      PAYMENT_PROCESSOR,
        }
        replacements = {k: v for k, v in replacements.items() if k and v}

        doc = Document(TEMPLATE_PATH)
        replace_in_document(doc, replacements)
        _fix_donor_section(doc, donor_email)
        _fix_date_font(doc)
        _fix_image_wrap(doc)

        # Negate the fee in table column 5 (fee column should show as -$x.xx)
        for table in doc.tables:
            for row in table.rows:
                cell = row.cells[5]
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.startswith("$"):
                            run.text = "-" + run.text

        temp_docx = os.path.join(OUTPUT_DIR, f"_temp_{doc_counter}.docx")
        base_name = safe_filename(donor_name, date_fmt, doc_id)
        # Deduplicate: same donor donating multiple times on the same day
        seen_filenames[base_name] = seen_filenames.get(base_name, 0) + 1
        if seen_filenames[base_name] > 1:
            stem, ext = base_name.rsplit(".", 1)
            base_name = f"{stem}_{seen_filenames[base_name]}.{ext}"
        output_pdf = os.path.join(OUTPUT_DIR, base_name)
        doc.save(temp_docx)

        import subprocess
        result = subprocess.run(
            [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "--headless", "--convert-to", "pdf",
                "--outdir", OUTPUT_DIR,
                temp_docx,
            ],
            capture_output=True, text=True,
        )
        os.remove(temp_docx)
        # LibreOffice names the output after the input file; rename to our target
        lo_output = os.path.join(OUTPUT_DIR, os.path.basename(temp_docx).replace(".docx", ".pdf"))
        if os.path.exists(lo_output):
            os.rename(lo_output, output_pdf)
        else:
            print(f"  WARNING: PDF not created for row {i}. soffice stderr: {result.stderr.strip()}")

        manifest.append({
            "donor_name":  donor_name,
            "donor_email": donor_email,
            "pdf_file":    os.path.basename(output_pdf),
        })
        print(f"  [{i}/{len(rows)}] {os.path.basename(output_pdf)}")

    # Write manifest so draft_emails.py knows which PDF belongs to whom
    manifest_path = os.path.join(OUTPUT_DIR, "manifest.csv")
    with open(manifest_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["donor_name", "donor_email", "pdf_file"])
        writer.writeheader()
        writer.writerows(manifest)
    print(f"\nDone. Manifest written to '{manifest_path}'.")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python generate_receipts.py <donations.csv>")
        sys.exit(1)
    generate_receipts(sys.argv[1])
